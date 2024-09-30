from django.shortcuts import render

# Create your views here.

from django.shortcuts import render
from order.models import Order, OrderDetails
from django.db.models import Sum
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from django.http import HttpResponse



def report_view(request):
    return render(request, 'dashboard/report_home.html')

# Báo cáo theo ngày
def daily_report(request):
    if request.method == 'POST':
        selected_date = request.POST.get('selected_date')
        if selected_date:
            selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date()

            # Lấy các đơn hàng trong ngày đã chọn
            orders = Order.objects.filter(order_date__date=selected_date)

            # Tính tổng doanh thu trong ngày
            total_revenue = orders.aggregate(Sum('total_price'))['total_price__sum'] or 0

            # Lấy thông tin các mặt hàng đã bán
            order_details = OrderDetails.objects.filter(order_id__in=orders)
            products_sold = order_details.values('dish_id__name').annotate(
                total_quantity=Sum('quantity'),
                unit_price=Sum('unit_price')
            )

            context = {
                'total_revenue': total_revenue,
                'products_sold': products_sold,
                'selected_date': selected_date
            }
            return render(request, 'dashboard/daily_report.html', context)

    return render(request, 'dashboard/daily_report.html')


def get_week_range(week_str=None):
    if not week_str:
        # Nếu không có tuần nào được chọn, dùng tuần hiện tại
        current_date = datetime.now()
        year, week = current_date.isocalendar()[0], current_date.isocalendar()[1]
    else:
        # Parse tuần từ định dạng "YYYY-WW"
        year, week = map(int, week_str.split('-W'))

    # Tìm ngày đầu tuần (thứ 2) của tuần đó
    start_date = datetime.strptime(f'{year} {week} 1', "%Y %W %w")
    # Ngày cuối tuần (chủ nhật) là 6 ngày sau ngày đầu tuần
    end_date = start_date + timedelta(days=6)

    return start_date, end_date, f'{year}-W{str(week).zfill(2)}'


def weekly_report(request):
    week_str = request.GET.get('week', None)
    start_date, end_date, selected_week = get_week_range(week_str)

    # Lấy các Order trong khoảng thời gian của tuần được chọn
    orders = Order.objects.filter(order_date__range=[start_date, end_date])

    # Lấy chi tiết đơn hàng từ các order này
    weekly_data = []
    for order in orders:
        order_details = OrderDetails.objects.filter(order_id=order.id)
        for detail in order_details:
            weekly_data.append({
                'name': detail.dish_id.name,  # Assuming you have a field `name` in `Dish`
                'quantity': detail.quantity,
                'unit_price': detail.unit_price,
                'total': detail.quantity * detail.unit_price
            })

    context = {
        'start_date': start_date,
        'end_date': end_date,
        'week_str': selected_week,
        'weekly_data': weekly_data
    }

    return render(request, 'dashboard/weekly_report.html', context)


def monthly_report(request):
    # Lấy tháng và năm từ request
    month = request.GET.get('month', datetime.now().month)
    year = request.GET.get('year', datetime.now().year)

    # Chuyển đổi tháng và năm thành kiểu số nguyên
    month = int(month)
    year = int(year)

    # Tính khoảng thời gian báo cáo
    start_date = datetime(year, month, 1)
    end_date = datetime(year, month + 1, 1) if month < 12 else datetime(year + 1, 1, 1)

    # Lấy các Order trong tháng
    orders = Order.objects.filter(order_date__range=[start_date, end_date])

    # Tính toán tổng số lượng và doanh thu
    total_quantity = orders.aggregate(total=Sum('orderdetails__quantity'))['total'] or 0
    total_revenue = orders.aggregate(total=Sum('total_price'))['total'] or 0

    # Tạo danh sách tháng
    months = [i for i in range(1, 13)]

    return render(request, 'dashboard/monthly_report.html', {
        'month': month,
        'year': year,
        'orders': orders,
        'total_quantity': total_quantity,
        'total_revenue': total_revenue,
        'months': months,  # Truyền danh sách tháng vào context
    })


def yearly_report(request):
    # Lấy năm từ request
    year = request.GET.get('year', datetime.now().year)

    # Chuyển đổi năm thành kiểu số nguyên
    year = int(year)

    # Lấy các Order trong năm
    orders = Order.objects.filter(order_date__year=year)

    # Tính toán tổng số lượng và doanh thu
    total_quantity = orders.aggregate(total=Sum('orderdetails__quantity'))['total'] or 0
    total_revenue = orders.aggregate(total=Sum('total_price'))['total'] or 0

    return render(request, 'dashboard/yearly_report.html', {
        'year': year,
        'orders': orders,
        'total_quantity': total_quantity,
        'total_revenue': total_revenue,
    })


def export_report_doc(request, report_type):
    # Lấy thông tin ngày, tháng, năm từ request
    if report_type == 'daily':
        selected_date = request.GET.get('selected_date')
        try:
            # Đảm bảo định dạng ngày theo đúng định dạng '%Y-%m-%d'
            selected_date = datetime.strptime(selected_date, '%Y-%m-%d').date()
        except ValueError:
            return HttpResponse("Invalid date format.", status=400)

        start_date = selected_date
        end_date = start_date + timedelta(days=1)
    elif report_type == 'weekly':
        week_str = request.GET.get('week', None)
        if not week_str:
            return HttpResponse("Please provide a valid week.", status=400)

        # Sử dụng hàm get_week_range để lấy khoảng thời gian của tuần
        start_date, end_date, selected_week = get_week_range(week_str)
    elif report_type == 'monthly':
        month = int(request.GET.get('month', datetime.now().month))
        year = int(request.GET.get('year', datetime.now().year))
        start_date = datetime(year, month, 1)
        end_date = datetime(year, month + 1, 1) if month < 12 else datetime(year + 1, 1, 1)
    elif report_type == 'yearly':
        year = int(request.GET.get('year', datetime.now().year))
        start_date = datetime(year, 1, 1)
        end_date = datetime(year + 1, 1, 1)
    else:
        return HttpResponse("Invalid report type.", status=400)

    # Lấy các Order trong khoảng thời gian đã chọn
    orders = Order.objects.filter(order_date__range=[start_date, end_date])

    # Tạo file .docx
    document = Document()

    # Thêm tiêu đề công ty
    company_name = "Company Name"
    paragraph = document.add_paragraph(company_name)
    paragraph.runs[0].font.size = Pt(14)

    # Thêm ngày xuất báo cáo
    today = datetime.today().strftime('%d-%m-%Y')
    document.add_paragraph(f"Date printed: {today}").runs[0].font.size = Pt(10)

    # Thêm khoảng thời gian báo cáo
    if report_type == 'daily':
        document.add_paragraph(f"Report for day: {selected_date}").runs[0].font.size = Pt(10)
    elif report_type == 'weekly':
        document.add_paragraph(
            f"Report for week: {start_date.strftime('%d-%m-%Y')} to {end_date.strftime('%d-%m-%Y')}").runs[
            0].font.size = Pt(10)
    elif report_type == 'monthly':
        document.add_paragraph(f"Report for month: {start_date.strftime('%B %Y')}").runs[0].font.size = Pt(10)
    elif report_type == 'yearly':
        document.add_paragraph(f"Report for year: {year}").runs[0].font.size = Pt(10)

    # Tạo bảng thông tin báo cáo
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total'

    # Thêm dữ liệu từ các Order vào bảng
    for order in orders:
        order_details = OrderDetails.objects.filter(order_id=order.id)
        for detail in order_details:
            row_cells = table.add_row().cells
            row_cells[0].text = detail.dish_id.name  # Sử dụng dish_id để lấy tên món ăn
            row_cells[1].text = str(detail.quantity)
            row_cells[2].text = f"{detail.unit_price:.2f}"
            row_cells[3].text = f"{detail.quantity * detail.unit_price:.2f}"

    # Tạo response để trả file .docx về người dùng
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{report_type}_report.docx"'

    # Lưu file vào response
    document.save(response)
    return response

